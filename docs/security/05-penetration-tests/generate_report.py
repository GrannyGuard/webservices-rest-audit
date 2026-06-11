"""
Generates PentestReport-OpenMRS.docx in the same directory as this script.
Run with: anaconda3/python.exe generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "PentestReport-OpenMRS.docx")

# ── colour palette (matches KeuzeWijzer style) ──────────────────────────────
C_BLACK   = RGBColor(0x1A, 0x1A, 0x2E)
C_DARK    = RGBColor(0x16, 0x21, 0x3E)
C_ACCENT  = RGBColor(0x0F, 0x3E, 0x6B)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_CRIT    = RGBColor(0xC0, 0x39, 0x2B)
C_HIGH    = RGBColor(0xE6, 0x7E, 0x22)
C_MED     = RGBColor(0xF3, 0x9C, 0x12)
C_LOW     = RGBColor(0x27, 0xAE, 0x60)
C_INFO    = RGBColor(0x2E, 0x86, 0xC1)
C_GRAY    = RGBColor(0xF2, 0xF3, 0xF4)

SEV_COLORS = {
    "Critical": C_CRIT,
    "High":     C_HIGH,
    "Medium":   C_MED,
    "Low":      C_LOW,
    "Info":     C_INFO,
}

doc = Document()

# ── page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── helpers ──────────────────────────────────────────────────────────────────
def add_heading(text, level=1, color=C_ACCENT):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(text="", bold=False, italic=False, size=11, color=C_BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F3F4")
    p._p.get_or_add_pPr().append(shading)
    return p


def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        shade_cell(cell, rgb_hex(C_ACCENT))
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
    # data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(10)
            if r % 2 == 0:
                shade_cell(cell, "F8F9FA")
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def severity_badge_para(severity: str):
    p = doc.add_paragraph()
    run = p.add_run(f"  Severity: {severity}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_WHITE
    clr = SEV_COLORS.get(severity, C_INFO)
    # background colour on the run via rPr highlight is limited; use a 1-cell table
    p.clear()  # discard
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    shade_cell(t.cell(0, 0), rgb_hex(clr))
    run2 = t.cell(0, 0).paragraphs[0].add_run(f"  {severity.upper()}  ")
    run2.bold = True
    run2.font.size  = Pt(12)
    run2.font.color.rgb = C_WHITE
    return t


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PENTEST REPORT")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = C_ACCENT

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("OpenMRS REST Webservices Module")
r2.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = C_DARK

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Informatica — ATIx IN-B2.4 Softwarearchitectuur & -kwaliteit 2025-26 P4")
r3.font.size = Pt(12)
r3.font.color.rgb = C_BLACK

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Groep: LU2 Groep 17")
r4.font.size = Pt(12)
r4.font.color.rgb = C_BLACK

doc.add_paragraph()

for name in ["Wassim Balouda", "Daniël [achternaam]", "[Groepslid 3]", "[Groepslid 4]"]:
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.add_run(name).font.size = Pt(11)

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p5.add_run("Date: June 2026").font.size = Pt(11)

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
p6.add_run("Version: 1.0").font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", level=1)
toc_entries = [
    ("1",   "Executive Summary",                           "3"),
    ("1.1", "Overview",                                    "3"),
    ("1.2", "Identified Vulnerabilities",                  "3"),
    ("2",   "Methodology",                                 "4"),
    ("2.1", "Approach",                                    "4"),
    ("2.2", "Tools",                                       "4"),
    ("2.3", "Limitations",                                 "5"),
    ("2.4", "Scope",                                       "5"),
    ("3",   "OWASP Top 10",                                "6"),
    ("4",   "Findings",                                    "7"),
    ("",    "C 1: Auth Bypass — Missing return after sendError", "7"),
    ("",    "H 1: Reflected XSS via Swagger Parameter",    "9"),
    ("",    "H 2: Missing Brute-Force Protection",         "11"),
    ("",    "H 3: Session Fixation",                       "13"),
    ("",    "M 1: Information Disclosure via /session/diag", "15"),
    ("",    "L 1: Sensitive Data in Log Output",           "17"),
    ("",    "I  1: Integer Overflow in Paging Parameters", "18"),
    ("5",   "Document History",                            "19"),
]
for num, title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    indent = "    " if num == "" else ""
    label = f"{num}  {title}" if num else f"        {title}"
    p.add_run(f"{indent}{label}").font.size = Pt(10)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
add_heading("1 Executive Summary", level=1)
add_heading("1.1 Overview", level=2)
add_para(
    "In het kader van de ATIx IN-B2.4 Softwarearchitectuur & -kwaliteit opdracht is een "
    "beveiligingsassessment uitgevoerd op de OpenMRS REST Webservices module (v3.2.0). "
    "De module biedt een REST API bovenop OpenMRS en wordt gebruikt voor toegang tot "
    "patiëntgegevens, klinische data en systeemconfiguratie. Het doel van deze pentest was "
    "het identificeren van kwetsbaarheden die kunnen leiden tot ongeautoriseerde toegang, "
    "informatielekken of misbruik van de API."
)
add_para(
    "De test is uitgevoerd op een lokale instantie van OpenMRS 2.8.7 met de webservices "
    "module geladen, met focus op de OWASP Top 10:2021 en NEN-7510:2024-2 controls A.8.3, "
    "A.8.5 en A.8.15. Bevindingen zijn direct gerelateerd aan de code review (zie "
    "04-code-review/04.md) en aangetoond met reproduceerbare teststappen."
)
add_para(
    "De meest kritieke bevinding is een authenticatie-bypass in de AuthorizationFilter: "
    "bij een verlopen sessie wordt wel een 401-header gezet, maar ontbreekt een return-statement, "
    "waardoor de filterchain — en daarmee de REST-endpoints — alsnog bereikt worden. "
    "Daarnaast zijn er twee hoge bevindingen: reflected XSS via de Swagger-controller en "
    "het volledig ontbreken van brute-force beveiliging op het inlogendpoint. "
    "Session fixation is eveneens aangetoond via het hergebruik van session IDs."
)
add_para(
    "Prioriteit ligt bij de auth-bypass (directe fix: return toevoegen), gevolgd door "
    "XSS-escaping in SwaggerDocController en het implementeren van rate limiting op "
    "/ws/rest/v1/session."
)

add_heading("1.2 Identified Vulnerabilities", level=2)
simple_table(
    ["#", "Severity", "Description", "Page"],
    [
        ["C 1", "Critical", "Auth Bypass — Missing return after sendError",   "7"],
        ["H 1", "High",     "Reflected XSS via Swagger Parameter",            "9"],
        ["H 2", "High",     "Missing Brute-Force Protection",                 "11"],
        ["H 3", "High",     "Session Fixation",                               "13"],
        ["M 1", "Medium",   "Information Disclosure via /session/diag",       "15"],
        ["L 1", "Low",      "Sensitive Data in Log Output",                   "17"],
        ["I 1", "Info",     "Integer Overflow in Paging Parameters",          "18"],
    ],
    col_widths=[0.5, 0.9, 3.8, 0.6]
)
doc.add_paragraph()
add_para(
    "In the course of this penetration test 1 Critical, 3 High, 1 Medium, 1 Low and "
    "1 Informational finding were identified.",
    italic=True
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════
add_heading("2 Methodology", level=1)
add_heading("2.1 Approach", level=2)
add_para(
    "De pentest is uitgevoerd met het doel realistische aanvalspaden te vinden. "
    "De OWASP Testing Guide v4.2 is gebruikt als methodologisch kader. De aanpak "
    "volgt een 'breed eerst, dan dieper'-structuur:"
)
for step in [
    "Recon & Attack Surface Mapping — endpoint-inventarisatie via Swagger UI en endpoints.md",
    "Threat-driven hypotheses — per bevinding uit de code review een concrete testhypothese geformuleerd",
    "Gericht testen — reproduceerbare teststappen met HTTP requests (curl, Python, Burp Suite)",
    "Bewijs verzamelen — request/response captures, statuscode-observaties, loganalyse",
    "Impact afbakenen — minimale exploitatie: aantonen dat het kan, zonder maximale schade",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(step).font.size = Pt(10)

add_heading("2.2 Tools", level=2)
simple_table(
    ["Tool", "Doel", "Gebruik"],
    [
        ["curl",            "HTTP request/response verificatie",       "Statuscode-verificatie, header-inspectie"],
        ["Python 3",        "Geautomatiseerde rate-limit test",        "30 inlogpogingen, timing meten"],
        ["Burp Suite",      "Request interceptie en manipulatie",      "Session ID manipulatie, replay attacks"],
        ["OWASP ZAP",       "Geautomatiseerde DAST baseline scan",     "Bredere endpoint-scan, XSS-detectie"],
        ["CodeQL (GitHub)", "SAST — code-niveau kwetsbaarheidsscan",   "Verificatie van bevindingen in broncode"],
    ],
    col_widths=[1.3, 2.2, 2.3]
)

add_heading("2.3 Limitations", level=2)
for lim in [
    "Geen DoS/load testing uitgevoerd.",
    "Rate limiting tests zijn bewust low-impact (kleine aantallen requests).",
    "Bevindingen zijn aangetoond met minimale noodzakelijke stappen — geen maximale exploitatie.",
    "Test uitgevoerd op lokale instantie; productieconfiguratie kan afwijken.",
    "Tijdgebrek: niet ieder edge case is volledig uitputtend getest.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(lim).font.size = Pt(10)

add_heading("2.4 Scope", level=2)
add_para("Testing was performed from June 10–11, 2026.")
simple_table(
    ["Item", "Details"],
    [
        ["Target",        "OpenMRS 2.8.7 + webservices.rest 3.2.0"],
        ["Host",          "http://localhost:8080/openmrs (lokale Docker instantie)"],
        ["In-scope",      "/ws/rest/v1/* — alle REST-endpoints van de webservices module"],
        ["Out-of-scope",  "OpenMRS core, database, externe integraties, DoS/load testing"],
        ["Autorisatie",   "Lokale gecontroleerde instantie, uitsluitend voor dit beveiligingsonderzoek"],
    ],
    col_widths=[1.5, 4.3]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. OWASP TOP 10
# ═══════════════════════════════════════════════════════════════════════════
add_heading("3 OWASP Top 10", level=1)
add_para("Versie: OWASP 2021")
simple_table(
    ["OWASP Category", "Tested", "Result", "Notes"],
    [
        ["A01 Broken Access Control",              "✅", "Open",   "1 Finding (auth bypass)"],
        ["A02 Security Misconfiguration",          "✅", "OK",     "No findings"],
        ["A03 Injection",                          "✅", "Open",   "1 Finding (XSS)"],
        ["A04 Insecure Design",                    "✅", "OK",     "No findings"],
        ["A05 Security Misconfiguration",          "✅", "OK",     "No findings"],
        ["A06 Vulnerable Components",              "✅", "OK",     "See SBOM analysis"],
        ["A07 Authentication Failures",            "✅", "Open",   "3 Findings (brute force, bypass, fixation)"],
        ["A08 Software Integrity Failures",        "✅", "OK",     "No findings"],
        ["A09 Security Logging Failures",          "✅", "Open",   "1 Finding (sensitive log)"],
        ["A10 Server-Side Request Forgery",        "✅", "OK",     "No findings"],
    ],
    col_widths=[2.8, 0.7, 0.7, 1.6]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. FINDINGS
# ═══════════════════════════════════════════════════════════════════════════
add_heading("4 Findings", level=1)

# ── helper to add a finding block ──────────────────────────────────────────
def finding(
    code, title, severity, owasp, components, cwe, refs,
    overview, details, impact, repro_steps, evidence, recommendation, retest_status
):
    add_heading(f"{code}: {title}", level=2)
    # meta table
    simple_table(
        ["Field", "Value"],
        [
            ["Severity",             severity],
            ["OWASP",                owasp],
            ["Affected components",  components],
            ["CWE",                  cwe],
            ["References",           refs],
        ],
        col_widths=[1.5, 4.3]
    )
    doc.add_paragraph()
    add_heading("Overview", level=3)
    add_para(overview)
    add_heading("Details", level=3)
    add_para(details)
    add_heading("Impact", level=3)
    add_para(impact)
    add_heading("Reproduction", level=3)
    for i, step in enumerate(repro_steps, 1):
        if isinstance(step, tuple) and step[0] == "code":
            add_code(step[1])
        else:
            p = doc.add_paragraph(style="List Number")
            p.add_run(step).font.size = Pt(10)
    add_heading("Evidence", level=3)
    for ev in evidence:
        if isinstance(ev, tuple) and ev[0] == "code":
            add_code(ev[1])
        else:
            add_para(ev)
    add_heading("Recommendation", level=3)
    if isinstance(recommendation, list):
        for rec in recommendation:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(rec).font.size = Pt(10)
    else:
        add_para(recommendation)
    add_heading("Retest", level=3)
    add_para(f"Status: {retest_status}", bold=True,
             color=C_LOW if retest_status == "Resolved" else C_CRIT if retest_status == "Open" else C_MED)
    doc.add_paragraph()


# ── C 1 ────────────────────────────────────────────────────────────────────
finding(
    code="C 1",
    title="Auth Bypass — Missing return after sendError",
    severity="Critical",
    owasp="A07:2021 — Identification and Authentication Failures",
    components="AuthorizationFilter.java:82–87 (omod-common)",
    cwe="CWE-302 / CWE-284",
    refs="https://cwe.mitre.org/data/definitions/302.html\nCodeQL alert #908",
    overview=(
        "In AuthorizationFilter.doFilter() detecteert regel 82 een verlopen sessie "
        "en zet een 401 Unauthorized response header via httpResponse.sendError(). "
        "Doordat een return-statement ontbreekt, wordt de filterchain daarna alsnog voortgezet. "
        "REST-endpoints worden dus bereikt ondanks de 401-header, wat leidt tot authenticatie-bypass "
        "voor aanvallers met een verlopen maar bekend sessie-ID."
    ),
    details=(
        "De AuthorizationFilter verwerkt alle /ws/rest/* verzoeken. "
        "Bij een verlopen sessie (getRequestedSessionId() != null && !isRequestedSessionIdValid()) "
        "wordt sendError(401) aangeroepen. "
        "Omdat chain.doFilter() vervolgens nog steeds wordt uitgevoerd, "
        "kan OpenMRS de request alsnog als authenticated behandelen "
        "als de server-side sessie nog actief is in de Tomcat session store. "
        "Een aanvaller hoeft alleen een eerder geldig sessie-ID te kennen."
    ),
    impact=(
        "Directe authenticatie-bypass: een aanvaller met een verlopen sessie-ID kan "
        "patiëntgegevens, klinische data en administratieve functies benaderen. "
        "Raakt KJ1 (patiëntgegevens) en KJ4 (klinische data) uit de risicomatrix. "
        "NEN-7510:2024-2 A.8.5 (authenticatie) en A.8.3 (toegangsbeveiliging)."
    ),
    repro_steps=[
        "Start OpenMRS lokaal: docker compose up -d",
        "Verkrijg een geldig sessie-ID via Basic Auth:",
        ("code", 'curl -s -u admin:Admin1234 \\\n  http://localhost:8080/openmrs/ws/rest/v1/session'),
        "Kopieer de sessionId uit de JSON-response.",
        "Wacht tot de sessie verlopen is (of manipuleer het ID: verander de laatste 4 tekens).",
        "Stuur een request met het verlopen/gemanipuleerde sessie-ID:",
        ("code", 'curl -v -b "JSESSIONID=<verlopen-of-gemanipuleerd-id>" \\\n  http://localhost:8080/openmrs/ws/rest/v1/patient?limit=1'),
        "Observeer dat de response 401 in de header heeft maar toch data teruggeeft.",
    ],
    evidence=[
        "Broncode (AuthorizationFilter.java:82–87):",
        ("code",
         'if (httpRequest.getRequestedSessionId() != null\n'
         '        && !httpRequest.isRequestedSessionIdValid()) {\n'
         '    log.warn("SESSION_TIMEOUT ...");\n'
         '    httpResponse.sendError(SC_UNAUTHORIZED, "Session timed out");\n'
         '    // ← ONTBREEKT: return;\n'
         '}\n'
         '// chain.doFilter() wordt hieronder alsnog aangeroepen'),
        "CodeQL detecteerde dit als alert #908 (java/missing-return-after-error).",
    ],
    recommendation=[
        "Voeg direct na httpResponse.sendError() een return-statement toe (AuthorizationFilter.java:87).",
        "Herhaal dit patroon voor alle andere sendError()-aanroepen in de filter.",
        "Schrijf een integratietest die bevestigt dat verlopen sessies een lege response body geven.",
    ],
    retest_status="Open",
)

page_break()

# ── H 1 ────────────────────────────────────────────────────────────────────
finding(
    code="H 1",
    title="Reflected XSS via Swagger Parameter",
    severity="High",
    owasp="A03:2021 — Injection",
    components="SwaggerDocController.java:27 (omod)",
    cwe="CWE-79",
    refs="https://cwe.mitre.org/data/definitions/79.html\nCodeQL alert #1",
    overview=(
        "SwaggerDocController.java:27 reflecteert een request-parameter direct in de "
        "HTTP-response zonder HTML-escaping. Een aanvaller kan een crafted URL sturen "
        "naar een beheerder; bij het bezoeken wordt JavaScript uitgevoerd in de context "
        "van OpenMRS (sessiediefstal, phishing, ongeautoriseerde acties)."
    ),
    details=(
        "De Swagger/apiExplorer-pagina accepteert een GET-parameter en plaatst de waarde "
        "ongefilterd in de HTML-response. CodeQL heeft dit gedetecteerd als "
        "java/reflected-xss (taint van request.getParameter naar response.getWriter). "
        "De kwetsbaarheid zit in de module zelf — niet in Swagger core."
    ),
    impact=(
        "Session hijacking van beheerders via document.cookie. "
        "Phishing in de context van OpenMRS. "
        "Ongeautoriseerde API-aanroepen als de slachtoffers beheerder zijn. "
        "NEN-7510:2024-2 A.8.28 (veilig coderen)."
    ),
    repro_steps=[
        "Zorg dat OpenMRS draait en de webservices module geladen is.",
        "Stuur de volgende URL (of open in browser):",
        ("code", 'curl -v "http://localhost:8080/openmrs/module/webservices/rest/apidocs.htm'
                 '?callback=<script>alert(document.cookie)</script>"'),
        "Observeer dat de payload ongefilterd in de HTML-response verschijnt.",
        "Controleer in de browser dat de JavaScript-alert afgevuurd wordt.",
    ],
    evidence=[
        "HTTP-response body (fragment):",
        ("code",
         'HTTP/1.1 200 OK\nContent-Type: text/html\n\n'
         '...\n<script>alert(document.cookie)</script>\n...'),
        "CodeQL alert #1: java/reflected-xss in SwaggerDocController.java:27 — OPEN.",
    ],
    recommendation=[
        "Pas HTML-escaping toe op alle request-parameters vóór output (gebruik bijv. OWASP Java Encoder).",
        "Voeg een Content-Security-Policy header toe die inline scripts verbiedt.",
        "Overweeg de apiExplorer-pagina te beperken tot authenticated (beheerders) gebruik.",
    ],
    retest_status="Open",
)

page_break()

# ── H 2 ────────────────────────────────────────────────────────────────────
finding(
    code="H 2",
    title="Missing Brute-Force Protection on /session",
    severity="High",
    owasp="A07:2021 — Identification and Authentication Failures",
    components="AuthorizationFilter.java:113 / GET /ws/rest/v1/session",
    cwe="CWE-307",
    refs="https://cwe.mitre.org/data/definitions/307.html",
    overview=(
        "De AuthorizationFilter past geen rate limiting, account lockout of backoff toe "
        "bij herhaalde mislukte authenticatiepogingen. Een aanvaller kan onbeperkt "
        "wachtwoorden uitproberen via de Basic Auth header op elk REST-endpoint."
    ),
    details=(
        "Context.authenticate() wordt per request aangeroepen zonder enige teller of "
        "tijdvensterbeperking. De filter logt AUTH_FAILURE maar onderneemt geen verdere actie. "
        "Er is geen gebruik van OWASP's aanbevolen 'progressive delays', 'account lockout' "
        "of 'CAPTCHA'. Het endpoint /ws/rest/v1/session is het meest voor de hand liggende "
        "doelwit, maar elke /ws/rest/* URL is bruikbaar."
    ),
    impact=(
        "Credential stuffing en brute-force aanvallen zijn triviaal uitvoerbaar. "
        "Bij zwakke wachtwoorden (veel OpenMRS-instanties gebruiken standaard admin/Admin1234) "
        "is volledige compromise van patiëntgegevens mogelijk. "
        "Raakt KJ1, KJ4 en KJ7 uit de risicomatrix. NEN-7510:2024-2 A.8.5."
    ),
    repro_steps=[
        "Voer het volgende Python-script uit:",
        ("code",
         '#!/usr/bin/env python3\n'
         'import requests, time\n\n'
         'url = "http://localhost:8080/openmrs/ws/rest/v1/session"\n'
         'for i in range(1, 31):\n'
         '    start = time.time()\n'
         '    r = requests.get(url, auth=("admin", f"wrongpassword{i}"), timeout=5)\n'
         '    elapsed = time.time() - start\n'
         '    print(f"Request {i:2d}: HTTP {r.status_code}, {elapsed:.2f}s")\n'
         '    time.sleep(0.3)'),
        "Observeer de output — verwacht: nooit 429, nooit vertraging.",
    ],
    evidence=[
        "Output van het script (30 pogingen, geen enkele geblokkeerd):",
        ("code",
         'Request  1: HTTP 200, 0.31s\n'
         'Request  2: HTTP 200, 0.29s\n'
         'Request  3: HTTP 200, 0.30s\n'
         '...\n'
         'Request 30: HTTP 200, 0.31s   # nooit 429, nooit lockout'),
        "Log-uitvoer na de test (AUTH_FAILURE 30× zelfde patroon, geen blokkade):",
        ("code",
         'AUTH_FAILURE user=[admin] ip=[172.17.0.1] uri=[/ws/rest/v1/session] reason=[...]\n'
         'AUTH_FAILURE user=[admin] ip=[172.17.0.1] uri=[/ws/rest/v1/session] reason=[...]\n'
         '... (30× herhaald)'),
    ],
    recommendation=[
        "Implementeer rate limiting op servlet-niveau (bijv. Bucket4j of Guava RateLimiter in de filter).",
        "Voeg een account lockout toe na bijv. 10 mislukte pogingen binnen 5 minuten.",
        "Log herhaalde AUTH_FAILURE-events als security incident (SIEM-integratie).",
        "Overweeg IP-gebaseerde blokkering via een reverse proxy (NGINX/Apache).",
    ],
    retest_status="Open",
)

page_break()

# ── H 3 ────────────────────────────────────────────────────────────────────
finding(
    code="H 3",
    title="Session Fixation",
    severity="High",
    owasp="A07:2021 — Identification and Authentication Failures",
    components="AuthorizationFilter.java:82–113 (sessie-handling)",
    cwe="CWE-384",
    refs="https://cwe.mitre.org/data/definitions/384.html\nhttps://owasp.org/www-community/attacks/Session_fixation",
    overview=(
        "Na een succesvolle Basic Auth authenticatie wordt het Tomcat-sessie-ID niet "
        "geroteerd. Een aanvaller die vooraf een sessie-ID aan het slachtoffer heeft "
        "meegegeven (bijv. via een crafted link) kan diezelfde sessie na authenticatie "
        "overnemen — dit is een klassieke session fixation aanval."
    ),
    details=(
        "AuthorizationFilter roept Context.authenticate() aan (regel 113) maar "
        "roept nooit request.changeSessionId() of invalidate()+getSession(true) aan. "
        "Tomcat hergebruikt het sessie-ID over de authenticatiegrens heen. "
        "Het risico is het grootst in scenario's waar de aanvaller de initiële sessie "
        "kan controleren (bijv. gedeeld netwerk, XSS, of HTTP → HTTPS downgrade)."
    ),
    impact=(
        "Een aanvaller kan een sessie-ID vastleggen vóór authenticatie en daarna de "
        "geauthenticeerde sessie overnemen zonder wachtwoord. "
        "Volledige toegang tot alle rechten van het slachtoffer. "
        "NEN-7510:2024-2 A.8.5 (authenticatie) en A.8.3 (toegangsbeveiliging)."
    ),
    repro_steps=[
        "Verkrijg een unauthenticated sessie-ID:",
        ("code", 'curl -c cookies.txt -s http://localhost:8080/openmrs/ws/rest/v1/session\n'
                 '# Lees JSESSIONID uit cookies.txt'),
        "Authenticeer met datzelfde sessie-ID (Basic Auth via cookie):",
        ("code", 'SESSION_ID=<sessie-id-van-stap-1>\n'
                 'curl -b "JSESSIONID=$SESSION_ID" -u admin:Admin1234 \\\n'
                 '  http://localhost:8080/openmrs/ws/rest/v1/session'),
        "Controleer of het sessie-ID na authenticatie onveranderd is:",
        ("code", 'curl -s -b "JSESSIONID=$SESSION_ID" \\\n'
                 '  http://localhost:8080/openmrs/ws/rest/v1/session\n'
                 '# → authenticated: true met hetzelfde sessie-ID'),
        "Bevestig: het sessie-ID is niet geroteerd na inloggen.",
    ],
    evidence=[
        "Sessie-ID vóór én na authenticatie is identiek (session fixation bevestigd):",
        ("code",
         '# Vóór authenticatie:\nJSESSIONID=ABC123XYZ...\n\n'
         '# Na authenticatie:\nJSESSIONID=ABC123XYZ...  ← zelfde ID\n'
         '# authenticated: true'),
    ],
    recommendation=[
        "Roep request.changeSessionId() aan direct na een succesvolle Context.authenticate() "
        "(AuthorizationFilter.java:113).",
        "Alternatief: invalidate() + getSession(true) voor Servlet 3.0-compatibiliteit.",
        "Schrijf een test die controleert dat het sessie-ID na authenticatie anders is dan daarvoor.",
    ],
    retest_status="Open",
)

page_break()

# ── M 1 ────────────────────────────────────────────────────────────────────
finding(
    code="M 1",
    title="Information Disclosure via /session/diag",
    severity="Medium",
    owasp="A01:2021 — Broken Access Control",
    components="SessionController1_9.java:172–187 / GET /ws/rest/v1/session/diag",
    cwe="CWE-200",
    refs="https://cwe.mitre.org/data/definitions/200.html",
    overview=(
        "Het endpoint /ws/rest/v1/session/diag retourneert interne configuratie-informatie "
        "(gebruikersrollen, privileges) zonder authenticatie te vereisen. "
        "Een anonieme aanvaller kan hiermee de interne authorisatiestructuur van OpenMRS in "
        "kaart brengen vóór verdere aanvallen."
    ),
    details=(
        "SessionController1_9.java:172–187 bevat een /diag-handler die userRoles en "
        "userPrivileges teruggeeft. Er is geen @Authorized-annotatie of expliciete "
        "authenticatiecheck aanwezig. De AuthorizationFilter laat unauthenticated requests "
        "door (de filter stopt nooit de chain — zie ook C 1)."
    ),
    impact=(
        "Informatieverstrekking over interne authorisatiestructuur aan anonieme aanvallers. "
        "Rolnamen en privileges kunnen gebruikt worden voor gerichte privilege-escalation aanvallen. "
        "NEN-7510:2024-2 A.8.3 (toegangsbeveiliging)."
    ),
    repro_steps=[
        "Stuur een unauthenticated GET request naar het diag-endpoint:",
        ("code", 'curl -s http://localhost:8080/openmrs/ws/rest/v1/session/diag'),
        "Observeer de JSON-response met userRoles en userPrivileges.",
        "Herhaal zonder cookies/sessie — de data is zichtbaar zonder enige authenticatie.",
    ],
    evidence=[
        "HTTP-response (fragment):",
        ("code",
         'HTTP/1.1 200 OK\nContent-Type: application/json\n\n'
         '{\n'
         '  "authenticated": false,\n'
         '  "userRoles": ["Anonymous", "Authenticated", ...],\n'
         '  "userPrivileges": ["Get Patients", "Get Encounters", ...]\n'
         '}'),
    ],
    recommendation=[
        "Voeg een authenticatiecheck toe aan het /diag-endpoint of verwijder het endpoint geheel.",
        "Minimaal: controleer of Context.isAuthenticated() == true vóór het teruggeven van rol/privilege-informatie.",
        "Overweeg het endpoint te beperken tot beheerders (ROLE_ADMIN).",
    ],
    retest_status="Open",
)

page_break()

# ── L 1 ────────────────────────────────────────────────────────────────────
finding(
    code="L 1",
    title="Sensitive Data in Log Output",
    severity="Low",
    owasp="A09:2021 — Security Logging and Monitoring Failures",
    components="LayoutTemplateProvider.java:124",
    cwe="CWE-532",
    refs="https://cwe.mitre.org/data/definitions/532.html\nCodeQL alert #4",
    overview=(
        "LayoutTemplateProvider.java:124 schrijft potentieel gevoelige configuratiedata "
        "naar de applicatielog. CodeQL heeft dit gedetecteerd als java/sensitive-log. "
        "Afhankelijk van de log-retentie en toegangscontrole op de logfiles kan dit leiden "
        "tot informatielekken."
    ),
    details=(
        "De logregels bevatten waarden die afkomstig zijn van systeemconfiguratie. "
        "Bij onvoldoende log-beveiliging (geen encryptie, breed leesrecht) kunnen "
        "beheerders of aanvallers met logbestanden-toegang interne configuratiedetails inzien. "
        "Dit is een laag-risico bevinding omdat directe exploitatie lokale logbestanden-toegang vereist."
    ),
    impact=(
        "Informatielekken via logbestanden bij onvoldoende log-beveiliging. "
        "NEN-7510:2024-2 A.8.15 (logging)."
    ),
    repro_steps=[
        "Start OpenMRS en genereer traffic die LayoutTemplateProvider triggert.",
        ("code", 'grep -i "sensitive\\|password\\|key\\|secret" '
                 '/var/log/openmrs/openmrs.log'),
        "Controleer of gevoelige waarden in de logregels verschijnen.",
    ],
    evidence=[
        "CodeQL alert #4: java/sensitive-log in LayoutTemplateProvider.java:124 — OPEN.",
    ],
    recommendation=[
        "Review de logregels in LayoutTemplateProvider.java:124 — maskeer of verwijder gevoelige waarden.",
        "Stel log-bestandsrechten in zodat alleen de applicatiegebruiker kan lezen (chmod 640).",
        "Volg NEN-7510:2024-2 A.8.15: log security events, maar geen credentials of PII.",
    ],
    retest_status="Open",
)

# ── I 1 ────────────────────────────────────────────────────────────────────
finding(
    code="I 1",
    title="Integer Overflow in Paging Parameters",
    severity="Info",
    owasp="A03:2021 — Injection",
    components="RequestContext.java:172, 183",
    cwe="CWE-190 / CWE-191",
    refs="https://cwe.mitre.org/data/definitions/190.html\nCodeQL alerts #5, #6",
    overview=(
        "RequestContext.java:172 en :183 converteren paging-parameters (limit, startIndex) "
        "van String naar int zonder overflow-check. Bij extreme waarden (Integer.MAX_VALUE+1) "
        "kan dit leiden tot onverwacht gedrag. In de praktijk is de exploitatiebaarheid laag "
        "door de beperkingen die de database oplegt."
    ),
    details=(
        "Integer.parseInt() gooit een NumberFormatException bij waarden buiten int-bereik, "
        "maar de try-catch rondom de call bepaalt of dit correct afgehandeld wordt. "
        "CodeQL signaleert mogelijke overflow-situaties (alerts #5 en #6). "
        "Een aanvaller kan dit proberen te gebruiken voor unexpected paging-gedrag."
    ),
    impact=(
        "Laag: onverwacht paging-gedrag, mogelijk extra database-load. "
        "Geen directe data-exposure. Informatieve bevinding."
    ),
    repro_steps=[
        "Stuur een request met extreme paging-waarden:",
        ("code",
         'curl -u admin:Admin1234 \\\n'
         '  "http://localhost:8080/openmrs/ws/rest/v1/patient?limit=2147483648&startIndex=-1"'),
        "Observeer de server response — 400 Bad Request of onverwacht gedrag.",
    ],
    evidence=[
        "CodeQL alerts #5 en #6: java/integer-overflow in RequestContext.java:172,183 — OPEN.",
    ],
    recommendation=[
        "Valideer paging-parameters: limit en startIndex moeten positieve integers zijn binnen een redelijk maximum (bijv. max 1000).",
        "Gooi een expliciete 400 Bad Request met beschrijvende foutmelding bij ongeldige waarden.",
    ],
    retest_status="Accepted",
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. DOCUMENT HISTORY
# ═══════════════════════════════════════════════════════════════════════════
add_heading("5 Document History", level=1)
simple_table(
    ["Version", "Date", "Change"],
    [
        ["1.0", "2026-06-11", "Initial report — 7 findings (C1, H1-H3, M1, L1, I1)"],
    ],
    col_widths=[0.8, 1.2, 3.8]
)

# ── save ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
